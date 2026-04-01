from docx import Document
import streamlit as st
import io

class DocumentGenerator:
    """Generate documents from markdown content"""
    
    @staticmethod
    def create_docx_from_markdown(markdown_content: str) -> bytes:
        """Create DOCX file from markdown content"""
        try:
            doc = Document()
            
            # Parse markdown and add to document
            lines = markdown_content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    # Add as paragraph with basic formatting
                    paragraph = doc.add_paragraph(line)
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            return doc_bytes.getvalue()
        except Exception as e:
            st.error(f"Error creating DOCX: {str(e)}")
            return None
